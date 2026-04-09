#!/usr/bin/env python3
"""
PRD Word 文档生成脚本

用法：
    echo '<JSON 数据>' | python generate_prd_docx.py --output output.docx
    python generate_prd_docx.py --input data.json --output output.docx

输入：JSON 格式的 PRD 数据（通过 stdin 或 --input 文件）
输出：专业排版的 .docx 文件
"""

import json
import sys
import argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("错误：python-docx 未安装。请运行：python -m pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── 样式常量 ──────────────────────────────────────────────

COLORS = {
    "primary": RGBColor(0x1B, 0x3A, 0x5C),      # 深蓝
    "secondary": RGBColor(0x2E, 0x86, 0xC1),     # 中蓝
    "accent": RGBColor(0x1A, 0xBC, 0x9C),        # 绿色
    "text": RGBColor(0x2C, 0x3E, 0x50),           # 深灰
    "light_text": RGBColor(0x7F, 0x8C, 0x8D),    # 浅灰
    "warning": RGBColor(0xE7, 0x4C, 0x3C),        # 红色
    "bg_header": "1B3A5C",                         # 表头背景
    "bg_alt_row": "F2F8FC",                        # 交替行背景
}

FONT_NAMES = {
    "heading": "PingFang SC",
    "body": "PingFang SC",
    "fallback_heading": "Microsoft YaHei",
    "fallback_body": "Microsoft YaHei",
}


def set_chinese_font(run, font_name=None, size=None, bold=False, color=None):
    """设置中文字体（同时设 eastAsia 属性）"""
    if font_name is None:
        font_name = FONT_NAMES["body"]
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color


def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_chinese_font(run, FONT_NAMES["heading"], bold=True, color=COLORS["primary"])
    return heading


def add_paragraph_styled(doc, text, bold=False, color=None, size=10.5):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, size=size, bold=bold, color=color or COLORS["text"])
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    return p


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_from_data(doc, headers, rows, col_widths=None):
    """创建带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_chinese_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, COLORS["bg_header"])
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_chinese_font(run, size=9.5, color=COLORS["text"])
            if row_idx % 2 == 1:
                set_cell_shading(cell, COLORS["bg_alt_row"])

    # 设置列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 表后间距
    return table


def add_source_note(doc, source_url):
    """添加来源标注"""
    if source_url:
        p = doc.add_paragraph()
        run = p.add_run(f"📌 来源：{source_url}")
        set_chinese_font(run, size=8, color=COLORS["light_text"])
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)


def add_missing_info_note(doc, field_name):
    """添加信息缺失标注"""
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️ {field_name}：未找到公开信息")
    set_chinese_font(run, size=9.5, color=COLORS["warning"])


def generate_prd(data, output_path):
    """生成 PRD Word 文档"""
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    product_name = data.get("product_name", "未知产品")
    generated_date = data.get("generated_date", datetime.now().strftime("%Y-%m-%d"))
    source_url = data.get("source_url", "")
    research_coverage = data.get("research_coverage", "N/A")

    # ── 封面 ──
    doc.add_paragraph()
    title = doc.add_heading(f"{product_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_chinese_font(run, FONT_NAMES["heading"], size=28, bold=True, color=COLORS["primary"])

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("产品分析 PRD")
    set_chinese_font(run, size=16, color=COLORS["secondary"])

    doc.add_paragraph()

    meta_info = [
        f"生成日期：{generated_date}",
        f"数据来源：{source_url}",
        f"研究完成度：{research_coverage}",
    ]
    for info in meta_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(info)
        set_chinese_font(run, size=10, color=COLORS["light_text"])

    doc.add_page_break()

    # ── 目录占位 ──
    add_heading_styled(doc, "目录", level=1)
    add_paragraph_styled(doc, "[请在 Word 中插入自动目录：引用 → 目录]", color=COLORS["light_text"])
    doc.add_page_break()

    # ── 各章节 ──
    sections_data = data.get("sections", {})

    # 1. Executive Summary
    add_heading_styled(doc, "1. Executive Summary", level=1)
    exec_summary = sections_data.get("executive_summary", {})
    if exec_summary.get("description"):
        add_paragraph_styled(doc, exec_summary["description"])
    if exec_summary.get("positioning"):
        add_paragraph_styled(doc, f"产品定位：{exec_summary['positioning']}", bold=True)
    if exec_summary.get("value_proposition"):
        add_paragraph_styled(doc, f"核心价值主张：{exec_summary['value_proposition']}", bold=True)
    if exec_summary.get("target_market"):
        add_paragraph_styled(doc, f"目标市场：{exec_summary['target_market']}", bold=True)
    add_source_note(doc, exec_summary.get("source"))

    # 2. 产品概述
    add_heading_styled(doc, "2. 产品概述", level=1)
    overview = sections_data.get("product_overview", {})
    overview_rows = []
    fields = [
        ("产品名称", "name"), ("官网", "website"), ("所属行业", "industry"),
        ("创立时间", "founded"), ("公司/团队", "company"), ("融资情况", "funding"),
    ]
    for label, key in fields:
        value = overview.get(key, "未找到公开信息")
        source = overview.get(f"{key}_source", "")
        overview_rows.append([label, value, source or "—"])
    add_table_from_data(doc, ["项目", "内容", "来源"], overview_rows, col_widths=[3, 8, 5])

    # 3. 问题定义 & 目标用户
    add_heading_styled(doc, "3. 问题定义 & 目标用户", level=1)
    problem = sections_data.get("problem_and_users", {})

    add_heading_styled(doc, "3.1 产品解决的问题", level=2)
    if problem.get("core_problem"):
        add_paragraph_styled(doc, f"核心问题：{problem['core_problem']}", bold=True)
    if problem.get("impact"):
        add_paragraph_styled(doc, f"问题影响：{problem['impact']}")
    if problem.get("alternatives"):
        add_paragraph_styled(doc, f"现有替代方案：{problem['alternatives']}")
    add_source_note(doc, problem.get("source"))

    add_heading_styled(doc, "3.2 目标用户画像", level=2)
    users = problem.get("target_users", [])
    if users:
        user_rows = [[u.get("type", ""), u.get("description", ""), u.get("scenario", "")] for u in users]
        add_table_from_data(doc, ["用户类型", "描述", "使用场景"], user_rows, col_widths=[3, 6, 6])
    else:
        add_missing_info_note(doc, "目标用户画像")

    add_heading_styled(doc, "3.3 Jobs to be Done", level=2)
    jtbd = problem.get("jtbd", {})
    if jtbd:
        add_paragraph_styled(doc, f"When {jtbd.get('when', '[未知]')}")
        add_paragraph_styled(doc, f"I want to {jtbd.get('want', '[未知]')}")
        add_paragraph_styled(doc, f"So I can {jtbd.get('so_that', '[未知]')}")
    else:
        add_missing_info_note(doc, "Jobs to be Done")

    # 4. 核心功能架构
    add_heading_styled(doc, "4. 核心功能架构", level=1)
    features = sections_data.get("features", {})
    feature_list = features.get("list", [])
    if feature_list:
        feature_rows = [
            [str(i + 1), f.get("name", ""), f.get("description", ""),
             f.get("differentiation", ""), f.get("source", "")]
            for i, f in enumerate(feature_list)
        ]
        add_table_from_data(doc, ["#", "功能名称", "描述", "差异化", "来源"], feature_rows,
                           col_widths=[1, 3, 5, 2.5, 4])

    if features.get("unique_selling_points"):
        add_heading_styled(doc, "差异化功能分析", level=2)
        for usp in features["unique_selling_points"]:
            add_paragraph_styled(doc, f"• {usp.get('point', '')} — 来源：{usp.get('source', '未标注')}")

    # 5. 商业模式 & 定价策略
    add_heading_styled(doc, "5. 商业模式 & 定价策略", level=1)
    business = sections_data.get("business_model", {})
    if business.get("revenue_model"):
        add_paragraph_styled(doc, f"收入模式：{business['revenue_model']}", bold=True)
    if business.get("pricing_strategy"):
        add_paragraph_styled(doc, f"定价策略：{business['pricing_strategy']}", bold=True)
    add_source_note(doc, business.get("source"))

    pricing = business.get("pricing_tiers", [])
    if pricing:
        pricing_rows = [[p.get("tier", ""), p.get("price", ""), p.get("features", ""), p.get("target", "")]
                       for p in pricing]
        add_table_from_data(doc, ["套餐", "价格", "核心功能", "目标用户"], pricing_rows,
                           col_widths=[3, 3, 6, 4])
    elif not business.get("revenue_model"):
        add_missing_info_note(doc, "商业模式与定价")

    # 6. 竞品分析
    add_heading_styled(doc, "6. 竞品分析", level=1)
    competitors = sections_data.get("competitors", {})
    comp_list = competitors.get("list", [])
    if comp_list:
        comp_rows = [[c.get("name", ""), c.get("website", ""), c.get("diff", ""), c.get("source", "")]
                    for c in comp_list]
        add_table_from_data(doc, ["竞品名称", "官网", "核心差异", "来源"], comp_rows,
                           col_widths=[3, 5, 5, 3])
    else:
        add_missing_info_note(doc, "竞品信息")

    comparison = competitors.get("comparison_matrix", [])
    if comparison:
        add_heading_styled(doc, "竞品对比矩阵", level=2)
        if comparison:
            headers = list(comparison[0].keys())
            rows = [list(row.values()) for row in comparison]
            add_table_from_data(doc, headers, rows)

    # 7. 用户评价
    add_heading_styled(doc, "7. 用户评价 & 市场反馈", level=1)
    reviews = sections_data.get("user_reviews", {})
    positive = reviews.get("positive", [])
    negative = reviews.get("negative", [])
    if positive:
        add_heading_styled(doc, "正面评价", level=2)
        review_rows = [[r.get("source", ""), r.get("summary", ""), r.get("url", "")] for r in positive]
        add_table_from_data(doc, ["来源", "评价摘要", "URL"], review_rows, col_widths=[3, 8, 5])
    if negative:
        add_heading_styled(doc, "负面评价 / 常见抱怨", level=2)
        review_rows = [[r.get("source", ""), r.get("summary", ""), r.get("url", "")] for r in negative]
        add_table_from_data(doc, ["来源", "问题摘要", "URL"], review_rows, col_widths=[3, 8, 5])
    if reviews.get("overall"):
        add_heading_styled(doc, "市场口碑总结", level=2)
        add_paragraph_styled(doc, reviews["overall"])
    if not positive and not negative:
        add_missing_info_note(doc, "用户评价")

    # 8. SWOT 分析
    add_heading_styled(doc, "8. SWOT 分析", level=1)
    swot = sections_data.get("swot", {})
    swot_items = [
        ("Strengths 优势", swot.get("strengths", [])),
        ("Weaknesses 劣势", swot.get("weaknesses", [])),
        ("Opportunities 机会", swot.get("opportunities", [])),
        ("Threats 威胁", swot.get("threats", [])),
    ]
    for label, items in swot_items:
        add_heading_styled(doc, label, level=2)
        if items:
            for item in items:
                point = item if isinstance(item, str) else item.get("point", "")
                evidence = "" if isinstance(item, str) else item.get("evidence", "")
                text = f"• {point}"
                if evidence:
                    text += f" — 依据：{evidence}"
                add_paragraph_styled(doc, text)
        else:
            add_missing_info_note(doc, label)

    # 9. 技术栈推测
    add_heading_styled(doc, "9. 技术栈推测", level=1)
    tech = sections_data.get("tech_stack", [])
    if tech:
        tech_rows = [[t.get("category", ""), t.get("technology", ""),
                      t.get("confidence", ""), t.get("evidence", "")]
                    for t in tech]
        add_table_from_data(doc, ["类别", "技术", "置信度", "依据"], tech_rows, col_widths=[3, 4, 3, 6])
    else:
        add_missing_info_note(doc, "技术栈")

    # 10. 总结
    add_heading_styled(doc, "10. 总结 & 关键洞察", level=1)
    summary = sections_data.get("summary", {})
    if summary.get("core_value"):
        add_heading_styled(doc, "产品核心价值", level=2)
        add_paragraph_styled(doc, summary["core_value"])
    if summary.get("insights"):
        add_heading_styled(doc, "关键洞察", level=2)
        for i, insight in enumerate(summary["insights"], 1):
            point = insight if isinstance(insight, str) else insight.get("point", "")
            evidence = "" if isinstance(insight, str) else insight.get("evidence", "")
            text = f"{i}. {point}"
            if evidence:
                text += f" — 依据：{evidence}"
            add_paragraph_styled(doc, text)
    if summary.get("info_gaps"):
        add_heading_styled(doc, "信息缺口", level=2)
        for gap in summary["info_gaps"]:
            add_paragraph_styled(doc, f"☐ {gap}")

    # 数据来源汇总
    sources = data.get("sources", [])
    if sources:
        add_heading_styled(doc, "数据来源汇总", level=2)
        source_rows = [[str(i + 1), s.get("type", ""), s.get("url", ""), s.get("used_in", "")]
                      for i, s in enumerate(sources)]
        add_table_from_data(doc, ["#", "来源类型", "URL", "使用章节"], source_rows,
                           col_widths=[1, 3, 8, 4])

    # ── 页脚声明 ──
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        f"本文档由 AI 辅助生成于 {generated_date}，基于公开信息整理。\n"
        "所有信息均标注来源，推测性内容已明确标注。请以官方信息为准。"
    )
    set_chinese_font(run, size=8, color=COLORS["light_text"])

    # ── 保存 ──
    doc.save(output_path)
    print(f"✅ PRD 文档已生成：{output_path}")


def main():
    parser = argparse.ArgumentParser(description="PRD Word 文档生成器")
    parser.add_argument("--input", "-i", help="输入 JSON 文件路径（不指定则从 stdin 读取）")
    parser.add_argument("--output", "-o", required=True, help="输出 .docx 文件路径")
    args = parser.parse_args()

    # 读取 JSON 数据
    if args.input:
        with open(args.input, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = json.load(sys.stdin)

    # 确保输出目录存在
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    generate_prd(data, str(output_path))


if __name__ == "__main__":
    main()
